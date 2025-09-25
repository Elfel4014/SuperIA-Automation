from docx import Document
from docx.shared import Inches

class WordAutomator:
    def __init__(self):
        pass

    def create_document(self, filename="documento_superia.docx", title="Documento Gerado pela SuperIA", content="Este é um documento de exemplo gerado automaticamente pela SuperIA."):
        try:
            document = Document()
            document.add_heading(title, level=1)
            document.add_paragraph(content)
            document.save(filename)
            print(f"Documento Word \'{filename}\' criado com sucesso.")
            return True
        except Exception as e:
            print(f"Erro ao criar documento Word: {e}")
            return False

    def add_paragraph_to_document(self, filename, paragraph_text):
        try:
            document = Document(filename)
            document.add_paragraph(paragraph_text)
            document.save(filename)
            print(f"Parágrafo adicionado ao documento \'{filename}\' com sucesso.")
            return True
        except Exception as e:
            print(f"Erro ao adicionar parágrafo ao documento Word: {e}")
            return False

    def add_heading_to_document(self, filename, heading_text, level=2):
        try:
            document = Document(filename)
            document.add_heading(heading_text, level=level)
            document.save(filename)
            print(f"Título adicionado ao documento \'{filename}\' com sucesso.")
            return True
        except Exception as e:
            print(f"Erro ao adicionar título ao documento Word: {e}")
            return False

    def add_table_to_document(self, filename, headers, data):
        try:
            document = Document(filename)
            table = document.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            for row_data in data:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = str(cell_data)
            document.save(filename)
            print(f"Tabela adicionada ao documento \'{filename}\' com sucesso.")
            return True
        except Exception as e:
            print(f"Erro ao adicionar tabela ao documento Word: {e}")
            return False

if __name__ == "__main__":
    automator = WordAutomator()
    test_doc = "test_superia_word.docx"

    print("\n--- Criando documento Word ---")
    automator.create_document(test_doc, "Relatório de Teste SuperIA", "Este é o relatório inicial.")

    print("\n--- Adicionando conteúdo ---")
    automator.add_paragraph_to_document(test_doc, "Um novo parágrafo com informações adicionais.")
    automator.add_heading_to_document(test_doc, "Seção 2: Dados Importantes")

    headers = ["Item", "Quantidade", "Preço"]
    data = [
        ["Produto A", 10, 25.50],
        ["Produto B", 5, 100.00]
    ]
    automator.add_table_to_document(test_doc, headers, data)

    print(f"\nVerifique o arquivo \'{test_doc}\' para ver o resultado.")

